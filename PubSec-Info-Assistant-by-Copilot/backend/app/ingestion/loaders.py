"""Document loaders for different file types with metadata extraction."""
import logging
from abc import ABC, abstractmethod
from io import BytesIO
from typing import Protocol

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import Document, DocumentMetadata, DocumentType

logger = logging.getLogger(__name__)


class DocumentLoaderProtocol(Protocol):
    """Protocol for document loaders."""

    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load document from bytes."""
        ...

    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from document."""
        ...


class BaseDocumentLoader(ABC):
    """Base class for document loaders."""

    def __init__(self, doc_type: DocumentType) -> None:
        """Initialize loader with document type."""
        self.doc_type = doc_type

    @abstractmethod
    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load document from bytes."""
        pass

    @abstractmethod
    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from document."""
        pass


class HTMLLoader(BaseDocumentLoader):
    """Loader for HTML documents."""

    def __init__(self) -> None:
        """Initialize HTML loader."""
        super().__init__(DocumentType.HTML)

    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load HTML document."""
        try:
            html_str = content.decode("utf-8")
            soup = BeautifulSoup(html_str, "lxml")

            # Extract text content
            text = soup.get_text(separator="\n", strip=True)

            # Extract metadata
            metadata = self.extract_metadata(content)
            metadata.file_size_bytes = len(content)

            return Document(
                tenant_id=tenant_id,
                filename=filename,
                doc_type=self.doc_type,
                content=text,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load HTML document {filename}: {e}")
            raise ValueError(f"Invalid HTML document: {e}") from e

    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from HTML."""
        html_str = content.decode("utf-8")
        soup = BeautifulSoup(html_str, "lxml")

        metadata = DocumentMetadata()

        # Extract title
        title_tag = soup.find("title")
        if title_tag:
            metadata.title = title_tag.get_text(strip=True)

        # Extract meta tags
        author_meta = soup.find("meta", attrs={"name": "author"})
        if author_meta and author_meta.get("content"):
            metadata.author = str(author_meta.get("content"))

        # Try to extract language
        html_tag = soup.find("html")
        if html_tag and html_tag.get("lang"):
            metadata.language = str(html_tag.get("lang", "en"))

        return metadata


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF documents."""

    def __init__(self) -> None:
        """Initialize PDF loader."""
        super().__init__(DocumentType.PDF)

    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load PDF document."""
        try:
            pdf_file = BytesIO(content)
            reader = PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)

            # Extract metadata
            metadata = self.extract_metadata(content)
            metadata.file_size_bytes = len(content)
            metadata.page_count = len(reader.pages)

            return Document(
                tenant_id=tenant_id,
                filename=filename,
                doc_type=self.doc_type,
                content=full_text,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load PDF document {filename}: {e}")
            raise ValueError(f"Invalid PDF document: {e}") from e

    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from PDF."""
        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)

        metadata = DocumentMetadata()

        # Extract PDF metadata
        pdf_metadata = reader.metadata
        if pdf_metadata:
            if pdf_metadata.title:
                metadata.title = pdf_metadata.title
            if pdf_metadata.author:
                metadata.author = pdf_metadata.author
            if pdf_metadata.creation_date:
                metadata.created_at = pdf_metadata.creation_date

        return metadata


class DOCXLoader(BaseDocumentLoader):
    """Loader for DOCX documents."""

    def __init__(self) -> None:
        """Initialize DOCX loader."""
        super().__init__(DocumentType.DOCX)

    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load DOCX document."""
        try:
            docx_file = BytesIO(content)
            doc = DocxDocument(docx_file)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)

            # Extract metadata
            metadata = self.extract_metadata(content)
            metadata.file_size_bytes = len(content)

            return Document(
                tenant_id=tenant_id,
                filename=filename,
                doc_type=self.doc_type,
                content=full_text,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load DOCX document {filename}: {e}")
            raise ValueError(f"Invalid DOCX document: {e}") from e

    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from DOCX."""
        docx_file = BytesIO(content)
        doc = DocxDocument(docx_file)

        metadata = DocumentMetadata()

        # Extract core properties
        core_props = doc.core_properties
        if core_props.title:
            metadata.title = core_props.title
        if core_props.author:
            metadata.author = core_props.author
        if core_props.created:
            metadata.created_at = core_props.created

        return metadata


class TXTLoader(BaseDocumentLoader):
    """Loader for plain text documents."""

    def __init__(self) -> None:
        """Initialize TXT loader."""
        super().__init__(DocumentType.TXT)

    def load(self, content: bytes, filename: str, tenant_id: str) -> Document:
        """Load TXT document."""
        try:
            text = content.decode("utf-8")

            # Extract metadata (minimal for plain text)
            metadata = DocumentMetadata(file_size_bytes=len(content))

            return Document(
                tenant_id=tenant_id,
                filename=filename,
                doc_type=self.doc_type,
                content=text,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to load TXT document {filename}: {e}")
            raise ValueError(f"Invalid TXT document: {e}") from e

    def extract_metadata(self, content: bytes) -> DocumentMetadata:
        """Extract metadata from TXT (minimal)."""
        return DocumentMetadata(file_size_bytes=len(content))


class DocumentLoaderFactory:
    """Factory for creating document loaders."""

    _loaders: dict[DocumentType, BaseDocumentLoader] = {
        DocumentType.HTML: HTMLLoader(),
        DocumentType.PDF: PDFLoader(),
        DocumentType.DOCX: DOCXLoader(),
        DocumentType.TXT: TXTLoader(),
    }

    @classmethod
    def get_loader(cls, doc_type: DocumentType) -> BaseDocumentLoader:
        """Get loader for document type."""
        loader = cls._loaders.get(doc_type)
        if not loader:
            raise ValueError(f"No loader available for document type: {doc_type}")
        return loader

    @classmethod
    def get_loader_by_extension(cls, extension: str) -> BaseDocumentLoader:
        """Get loader by file extension."""
        extension = extension.lower().lstrip(".")
        doc_type_map = {
            "html": DocumentType.HTML,
            "htm": DocumentType.HTML,
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "txt": DocumentType.TXT,
        }
        doc_type = doc_type_map.get(extension)
        if not doc_type:
            raise ValueError(f"Unsupported file extension: {extension}")
        return cls.get_loader(doc_type)
